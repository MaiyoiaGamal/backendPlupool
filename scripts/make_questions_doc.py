#!/usr/bin/env python3
"""
Create a Word (.docx) file from one or more images.

Use-case: user has photos of homework/questions and wants a single Word file.

Example:
  python3 scripts/make_questions_doc.py --out "اسئلة.docx" الورقة1.jpg الورقة2.jpg
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Inches


def add_image_page(doc: Document, image_path: Path, index: int) -> None:
    doc.add_paragraph(f"الورقة رقم ({index})")
    # Keep a reasonable width so it fits on page
    doc.add_picture(str(image_path), width=Inches(6.0))
    doc.add_page_break()


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate a Word file containing the provided images."
    )
    parser.add_argument(
        "images",
        nargs="+",
        help="Paths to image files (jpg/png/webp/gif).",
    )
    parser.add_argument(
        "--out",
        default="اسئلة.docx",
        help="Output .docx path (default: اسئلة.docx).",
    )
    args = parser.parse_args()

    image_paths = [Path(p).expanduser().resolve() for p in args.images]
    for p in image_paths:
        if not p.exists():
            raise SystemExit(f"Image not found: {p}")

    out_path = Path(args.out).expanduser().resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_paragraph("الأسئلة (مجمّعة من الصور)")
    doc.add_paragraph("")

    for idx, img in enumerate(image_paths, start=1):
        add_image_page(doc, img, idx)

    doc.save(str(out_path))
    print(f"Created: {out_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())

